from __future__ import annotations

import os
import tempfile
import threading
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from src.core.services.knowledge import KnowledgeService
from src.core.services.memory import MemoryService, ModelMemoryExtractor
from src.core.plugins.todo import execute_action
from src.core.storage.tenants import ConversationStore, TenantRegistry
from src.core.modeling import (
    CanonicalMessage,
    ModelCapabilities,
    ModelIdentity,
    ModelResponse,
    ModelRouter,
    RerankError,
)


class FakeEmbedding:
    model_id = "fake-embedding"
    dimensions = 4

    def embed(self, texts):
        vectors = []
        for text in texts:
            if "水果" in text or "苹果" in text:
                vectors.append([1.0, 0.0, 0.0, 0.0])
            else:
                vectors.append([0.0, 1.0, 0.0, 0.0])
        return vectors

    def close(self):
        pass


class FakeRerank:
    model_id = "fake-rerank"

    def __init__(self, fail=False):
        self.fail = fail
        self.calls = []

    def rerank(self, query, documents, top_n=None):
        self.calls.append((query, list(documents), top_n))
        if self.fail:
            raise RerankError("rerank 服务不可用")
        # Reverse the candidate order with descending scores.
        indices = list(reversed(range(len(documents))))
        return [(index, float(len(indices) - rank)) for rank, index in enumerate(indices)]

    def close(self):
        pass


class FakeExtractor:
    def __init__(self, values):
        self.values = values

    def extract(self, _question, _answer):
        return list(self.values)

    def close(self):
        pass


class FlakyExtractor:
    def __init__(self):
        self.calls = 0

    def extract_with_status(self, _question, _answer):
        self.calls += 1
        if self.calls == 1:
            return False, []
        return True, [{
            "kind": "preference",
            "key": "reply-style",
            "content": "用户偏好简洁回答",
            "confidence": 0.99,
            "evidence_type": "explicit",
        }]

    def close(self):
        pass


class CompactingExtractor(FakeExtractor):
    def compact(self, items):
        return [{
            "kind": "identity",
            "content": "用户有多项经过归并的稳定背景信息",
            "source_memory_ids": [
                item["memory_id"] for item in items
            ],
        }]


class FakeDefaultModel:
    identity = ModelIdentity("default", "cloud", "default-model")
    capabilities = ModelCapabilities()

    def __init__(self, responses):
        self.responses = list(responses)
        self.calls = []

    def ensure_ready(self):
        pass

    def complete(self, request):
        self.calls.append(request)
        return ModelResponse(
            CanonicalMessage("assistant", self.responses.pop(0)),
            actual_model="default-model",
        )

    def close(self):
        pass


class SqliteKnowledgeMemoryTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.data = Path(self.temp.name) / "data"
        self.registry = TenantRegistry(self.data)
        self.tenant = self.registry.resolve("bot", "user")

    def test_database_pragmas_permissions_and_no_runtime_json(self):
        self.assertEqual(os.stat(self.registry.database_path).st_mode & 0o777, 0o600)
        with self.registry.database.read() as connection:
            self.assertEqual(connection.execute("PRAGMA foreign_keys").fetchone()[0], 1)
            self.assertEqual(connection.execute("PRAGMA journal_mode").fetchone()[0], "wal")
            self.assertEqual(connection.execute("PRAGMA busy_timeout").fetchone()[0], 5000)
        conversation = ConversationStore(self.registry, 12)
        conversation.append_transcript(self.tenant.tenant_id, "user", "只进入数据库")
        self.assertFalse(list(self.registry.tenant_root(self.tenant.tenant_id).rglob("*.json")))
        self.assertFalse(list(self.registry.tenant_root(self.tenant.tenant_id).rglob("*.jsonl")))

    def test_legacy_user_index_is_ignored(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory) / "data"
            (root / "system").mkdir(parents=True)
            (root / "system" / "users.json").write_text(
                '{"version":1,"users":{"old":{"tenant_id":"00000000-0000-0000-0000-000000000001"}}}',
                encoding="utf-8",
            )
            registry = TenantRegistry(root)
            self.assertEqual(registry.list_contexts(), [])

    def test_knowledge_idempotence_file_boundary_and_hybrid_search(self):
        service = KnowledgeService(self.registry, FakeEmbedding())
        first = service.add_text(
            self.tenant.tenant_id, "饮食", "苹果是一种常见水果，也可以用于制作甜点。"
        )
        again = service.add_text(
            self.tenant.tenant_id, "饮食", "苹果是一种常见水果，也可以用于制作甜点。"
        )
        self.assertFalse(first["unchanged"])
        self.assertTrue(again["unchanged"])
        self.assertEqual(len(service.list(self.tenant.tenant_id)), 1)
        results = service.search(self.tenant.tenant_id, "我想找水果资料")
        self.assertIn("苹果", results[0]["content"])

        workspace = self.registry.tenant_root(self.tenant.tenant_id) / "workspace"
        document = workspace / "notes.md"
        document.write_text("# 标题\n\nworkspace 知识", encoding="utf-8")
        indexed = service.index_file(self.tenant, document)
        self.assertEqual(indexed["status"], "ready")
        outside = Path(self.temp.name) / "outside.md"
        outside.write_text("outside", encoding="utf-8")
        with self.assertRaises(ValueError):
            service.index_file(self.tenant, outside)
        self.assertTrue(service.delete(self.tenant.tenant_id, indexed["source_id"]))

    def test_rich_document_indexing_boundary_and_limits(self):
        import docx
        from openpyxl import Workbook

        service = KnowledgeService(self.registry, None)
        workspace = self.registry.tenant_root(self.tenant.tenant_id) / "workspace"

        word_path = workspace / "manual.docx"
        document = docx.Document()
        document.add_heading("产品手册", level=1)
        document.add_paragraph("退货政策支持七天无理由退货。")
        document.save(str(word_path))
        indexed = service.index_file(self.tenant, word_path)
        self.assertEqual(indexed["status"], "pending_embedding")
        self.assertGreater(indexed["chunks"], 0)
        results = service.search(self.tenant.tenant_id, "退货政策")
        self.assertTrue(any("退货" in item["content"] for item in results))

        sheet_path = workspace / "price.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "价格表"
        sheet.append(["商品", "价格"])
        sheet.append(["苹果", 5])
        workbook.save(str(sheet_path))
        indexed = service.index_file(self.tenant, sheet_path)
        self.assertGreater(indexed["chunks"], 0)
        results = service.search(self.tenant.tenant_id, "价格表")
        self.assertTrue(any("苹果" in item["content"] for item in results))

        # Rich documents outside the tenant workspace stay rejected.
        outside = Path(self.temp.name) / "outside.docx"
        document.save(str(outside))
        with self.assertRaises(ValueError):
            service.index_file(self.tenant, outside)

        # The dedicated 20 MiB budget applies to rich documents.
        with patch("src.core.services.knowledge.MAX_DOCUMENT_BYTES", 10):
            with self.assertRaisesRegex(ValueError, "20 MiB"):
                service.index_file(self.tenant, word_path)

        unsupported = workspace / "data.csv"
        unsupported.write_text("a,b", encoding="utf-8")
        with self.assertRaises(ValueError):
            service.index_file(self.tenant, unsupported)

    def test_embedding_failure_keeps_searchable_chunks(self):
        service = KnowledgeService(self.registry, None)
        result = service.add_text(self.tenant.tenant_id, "离线", "中文全文检索仍然可用。")
        self.assertEqual(result["status"], "pending_embedding")
        self.assertEqual(len(service.search(self.tenant.tenant_id, "全文检索")), 1)

    def test_rerank_reorders_candidates_and_degrades_silently(self):
        docs = [
            ("苹果", "苹果是一种常见水果，可以直接食用。"),
            ("香蕉", "香蕉也是常见水果，富含钾元素。"),
            ("橙子", "橙子是水果，含有丰富维生素 C。"),
        ]
        plain = KnowledgeService(self.registry, FakeEmbedding())
        for name, text in docs:
            plain.add_text(self.tenant.tenant_id, name, text)
        baseline = [
            hit["content"]
            for hit in plain.search(self.tenant.tenant_id, "水果", limit=3)
        ]
        self.assertEqual(len(baseline), 3)

        reranker = FakeRerank()
        reranked_service = KnowledgeService(self.registry, FakeEmbedding(), reranker)
        reranked = [
            hit["content"]
            for hit in reranked_service.search(self.tenant.tenant_id, "水果", limit=3)
        ]
        self.assertEqual(reranked, list(reversed(baseline)))
        self.assertTrue(reranker.calls)
        self.assertEqual(reranker.calls[0][2], 3)

        degraded_service = KnowledgeService(
            self.registry, FakeEmbedding(), FakeRerank(fail=True)
        )
        degraded = [
            hit["content"]
            for hit in degraded_service.search(self.tenant.tenant_id, "水果", limit=3)
        ]
        self.assertEqual(degraded, baseline)

    def test_memory_review_conflict_secret_and_forget(self):
        extractor = FakeExtractor([
            {"kind": "preference", "key": "drink", "content": "用户喜欢喝绿茶", "confidence": 0.95},
            {"kind": "goal", "key": "career", "content": "用户想学习数据库", "confidence": 0.6},
            {"kind": "identity", "key": "secret", "content": "token: abc123", "confidence": 1.0},
        ])
        service = MemoryService(self.registry, extractor)
        self.addCleanup(service.close)
        ConversationStore(self.registry, 12).append_transcript(
            self.tenant.tenant_id, "user", "记住我的偏好"
        )
        created = service.extract(self.tenant.tenant_id, "记住我的偏好", "好的")
        self.assertEqual(len(created), 2)
        with self.registry.database.read() as connection:
            sources = connection.execute(
                "SELECT source_event_ids FROM memory_items WHERE memory_id=?", (created[0],)
            ).fetchone()[0]
        self.assertNotEqual(sources, "[]")
        items = service.list(self.tenant.tenant_id)
        self.assertEqual({item["status"] for item in items}, {"active", "pending"})
        pending = next(item for item in items if item["status"] == "pending")
        self.assertTrue(service.confirm(self.tenant.tenant_id, pending["memory_id"][:8]))
        active = service.search(self.tenant.tenant_id, "绿茶")
        tea = next(item for item in active if "绿茶" in item["content"])
        self.assertTrue(service.forget(self.tenant.tenant_id, tea["memory_id"][:8]))
        self.assertNotIn("绿茶", " ".join(item["content"] for item in service.search(self.tenant.tenant_id, "绿茶")))

    def test_model_memory_extractor_uses_default_model_router(self):
        default = FakeDefaultModel([
            '{"memories":[{"kind":"preference","key":"style",'
            '"content":"用户偏好简洁回答","confidence":0.99,'
            '"evidence_type":"explicit"}]}',
            '{"items":[{"kind":"preference","content":"用户偏好简洁回答",'
            '"source_memory_ids":["memory-1"]}]}',
        ])
        router = ModelRouter(
            {"default": default},
            primary_profile_id="default",
            fallback_profile_id="default",
        )
        extractor = ModelMemoryExtractor(router)

        succeeded, memories = extractor.extract_with_status("请简洁回答", "好的")
        compacted = extractor.compact([{
            "memory_id": "memory-1",
            "kind": "preference",
            "content": "用户偏好简洁回答",
        }])

        self.assertTrue(succeeded)
        self.assertEqual(memories[0]["content"], "用户偏好简洁回答")
        self.assertEqual(compacted[0]["source_memory_ids"], ["memory-1"])
        self.assertEqual(len(default.calls), 2)
        self.assertTrue(
            all(call.generation.reasoning is False for call in default.calls)
        )

    def test_soul_projection_is_isolated_private_and_rebuildable(self):
        other = self.registry.resolve("bot", "other")
        extractor = FakeExtractor([
            {
                "kind": "preference",
                "key": "reply-style",
                "content": "用户偏好简洁、直接的回答",
                "confidence": 0.99,
                "evidence_type": "explicit",
            },
            {
                "kind": "goal",
                "key": "career",
                "content": "用户可能想学习数据库",
                "confidence": 0.95,
                "evidence_type": "inferred",
            },
            {
                "kind": "identity",
                "key": "medical",
                "content": "用户的医疗记录包含高血压",
                "confidence": 1.0,
                "evidence_type": "explicit",
            },
        ])
        service = MemoryService(self.registry, extractor)
        self.addCleanup(service.close)
        created = service.extract(
            self.tenant.tenant_id,
            "请记住我的习惯",
            "好的",
            source_event_ids=[1],
        )
        self.assertEqual(len(created), 2)

        path = self.registry.tenant_root(self.tenant.tenant_id) / "SOUL.md"
        content = path.read_text(encoding="utf-8")
        self.assertIn("用户偏好简洁、直接的回答", content)
        self.assertNotIn("数据库", content)
        self.assertNotIn("医疗", content)
        self.assertLessEqual(len(content), 1200)
        if os.name != "nt":
            self.assertEqual(path.stat().st_mode & 0o777, 0o600)
        self.assertFalse(
            (self.registry.tenant_root(other.tenant_id) / "SOUL.md").exists()
        )

        with self.registry.database.transaction(immediate=True) as connection:
            connection.execute(
                "INSERT INTO memory_items("
                "memory_id, tenant_id, kind, content, normalized_key, confidence, "
                "status, source_event_ids, created_at, updated_at"
                ") VALUES (?, ?, 'identity', ?, 'identity:legacy-sensitive', 1, "
                "'active', '[]', ?, ?)",
                (
                    "legacy-sensitive",
                    self.tenant.tenant_id,
                    "用户患有糖尿病",
                    "2026-07-24T00:00:00+00:00",
                    "2026-07-24T00:00:00+00:00",
                ),
            )
        service.rebuild_soul(self.tenant.tenant_id)
        self.assertNotIn(
            "糖尿病",
            service.get_soul(self.tenant.tenant_id)["content"],
        )
        self.assertNotIn(
            "糖尿病",
            " ".join(
                item["content"]
                for item in service.search(self.tenant.tenant_id, "糖尿病")
            ),
        )

        profile = service.get_soul(self.tenant.tenant_id)
        self.assertEqual(
            service.search(
                self.tenant.tenant_id,
                "简洁",
                exclude_soul=True,
            ),
            [],
        )
        path.write_text("被人工篡改", encoding="utf-8")
        rebuilt = service.get_soul(self.tenant.tenant_id)
        self.assertNotEqual(rebuilt["content"], "被人工篡改")
        self.assertGreater(rebuilt["revision"], profile["revision"])

        pending = next(
            item for item in service.list(self.tenant.tenant_id)
            if item["status"] == "pending"
        )
        self.assertTrue(
            service.confirm(self.tenant.tenant_id, pending["memory_id"][:8])
        )
        self.assertIn(
            "数据库",
            service.get_soul(self.tenant.tenant_id)["content"],
        )

    def test_soul_projection_enforces_hard_budget_and_item_limit(self):
        extractor = FakeExtractor([
            {
                "kind": "identity",
                "key": "fact-{}".format(index),
                "content": "用户稳定背景信息第{}项：{}".format(index, "长内容" * 30),
                "confidence": 0.99,
                "evidence_type": "explicit",
            }
            for index in range(30)
        ])
        service = MemoryService(self.registry, extractor)
        self.addCleanup(service.close)
        service.extract(self.tenant.tenant_id, "记住这些背景", "好的")
        profile = service.get_soul(self.tenant.tenant_id)
        bullets = [
            line for line in profile["content"].splitlines() if line.startswith("- ")
        ]
        self.assertLessEqual(len(profile["content"]), 1200)
        self.assertLessEqual(len(bullets), 16)
        self.assertTrue(all(len(line[2:]) <= 80 for line in bullets))
        self.assertFalse(
            list(
                self.registry.tenant_root(self.tenant.tenant_id).glob(
                    ".SOUL.*.tmp"
                )
            )
        )

    def test_soul_over_soft_limit_uses_validated_local_compaction(self):
        extractor = CompactingExtractor([
            {
                "kind": "identity",
                "key": "compact-{}".format(index),
                "content": "用户稳定背景第{}项：{}".format(index, "背景内容" * 15),
                "confidence": 0.99,
                "evidence_type": "explicit",
            }
            for index in range(14)
        ])
        service = MemoryService(self.registry, extractor)
        self.addCleanup(service.close)
        service.extract(self.tenant.tenant_id, "记住这些长期背景", "好的")
        profile = service.get_soul(self.tenant.tenant_id)
        self.assertIn("经过归并", profile["content"])
        self.assertEqual(len(profile["source_memory_ids"]), 14)
        with self.registry.database.read() as connection:
            compacted_at = connection.execute(
                "SELECT compacted_at FROM soul_profiles WHERE tenant_id=?",
                (self.tenant.tenant_id,),
            ).fetchone()[0]
        self.assertIsNotNone(compacted_at)

    def test_daily_scan_retries_failed_local_extraction(self):
        conversation = ConversationStore(self.registry, 12)
        conversation.append_transcript(
            self.tenant.tenant_id, "user", "以后请简洁回答"
        )
        service = MemoryService(self.registry, FlakyExtractor())
        self.addCleanup(service.close)

        self.assertEqual(service.scan_tenant(self.tenant.tenant_id), 0)
        with self.registry.database.read() as connection:
            cursor = connection.execute(
                "SELECT last_scanned_event_id FROM soul_profiles WHERE tenant_id=?",
                (self.tenant.tenant_id,),
            ).fetchone()[0]
        self.assertEqual(cursor, 0)

        self.assertEqual(service.scan_tenant(self.tenant.tenant_id), 1)
        self.assertIn(
            "用户偏好简洁回答",
            service.get_soul(self.tenant.tenant_id)["content"],
        )

    def test_soul_atomic_write_never_exposes_partial_content(self):
        path = self.registry.tenant_root(self.tenant.tenant_id) / "SOUL.md"
        versions = {
            "<!-- auto-generated; revision: 1 -->\n" + "甲" * 800 + "\n",
            "<!-- auto-generated; revision: 2 -->\n" + "乙" * 800 + "\n",
        }
        MemoryService._atomic_write(path, next(iter(versions)))
        failures = []

        def writer():
            try:
                for _ in range(30):
                    for content in versions:
                        MemoryService._atomic_write(path, content)
            except Exception as exc:
                failures.append(exc)

        thread = threading.Thread(target=writer)
        thread.start()
        while thread.is_alive():
            self.assertIn(path.read_text(encoding="utf-8"), versions)
        thread.join()
        self.assertEqual(failures, [])
        self.assertIn(path.read_text(encoding="utf-8"), versions)

    def test_sqlite_todo_backend(self):
        result = execute_action(
            self.registry.database_path,
            self.tenant.tenant_id,
            "add",
            title="SQLite 待办",
        )
        self.assertIn("T0001", result.summary)
        listed = execute_action(
            self.registry.database_path,
            self.tenant.tenant_id,
            "list",
        )
        self.assertIn("SQLite 待办", listed.summary)
        self.assertFalse(
            (
                self.registry.tenant_root(self.tenant.tenant_id)
                / "scripts" / "todo" / "todos.json"
            ).exists()
        )

    def test_tenant_deletion_purges_fts_and_keeps_audit(self):
        service = KnowledgeService(self.registry, None)
        service.add_text(self.tenant.tenant_id, "private", "只能属于当前租户的秘密知识")
        tenant_id = self.tenant.tenant_id
        self.registry.delete(self.tenant)
        with self.registry.database.read() as connection:
            self.assertEqual(
                connection.execute(
                    "SELECT COUNT(*) FROM knowledge_fts WHERE tenant_id=?", (tenant_id,)
                ).fetchone()[0],
                0,
            )
            self.assertEqual(
                connection.execute(
                    "SELECT status FROM deletion_audit WHERE tenant_id=?", (tenant_id,)
                ).fetchone()[0],
                "complete",
            )


if __name__ == "__main__":
    unittest.main()
