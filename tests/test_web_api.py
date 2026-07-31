from __future__ import annotations

import json
import os
import tempfile
import threading
import unittest
from dataclasses import replace
from unittest.mock import MagicMock, patch

from bs4 import BeautifulSoup
from fastapi.testclient import TestClient

from src.core.modeling import (
    CanonicalMessage,
    CanonicalToolCall,
    ModelCapabilities,
    ModelCallContext,
    ModelIdentity,
    ModelResponse,
    ModelRouter,
)
from src.core.tooling.models import ToolResult


class FakeClient:
    def __init__(self, profile_id="test_model", provider="test", model="test-model"):
        self.identity = ModelIdentity(profile_id, provider, model)
        self.capabilities = ModelCapabilities(tools=True, vision=False, reasoning=False)
        self.closed = False

    def ensure_ready(self):
        pass

    def complete(self, request):
        return ModelResponse(
            CanonicalMessage("assistant", "你好！我是测试模型。"),
            actual_model=self.identity.configured_model,
        )

    def complete_stream(self, request):
        yield "你好！"
        yield "我是测试模型。"

    def close(self):
        self.closed = True


class MultiAgentToolClient(FakeClient):
    """Return one structured directory call per agent, then a final answer."""

    def __init__(self):
        super().__init__()
        self.requests = []
        self._requests_lock = threading.Lock()

    def complete(self, request):
        with self._requests_lock:
            self.requests.append(request)
        agent_id = request.context.agent_id or "unknown"
        tool_messages = [message for message in request.messages if message.role == "tool"]
        if tool_messages:
            payload = json.loads(tool_messages[-1].content)
            if payload.get("ok"):
                text = "{}：目录读取完成。".format(agent_id)
            else:
                text = "{}：目录读取失败。".format(agent_id)
            return ModelResponse(
                CanonicalMessage("assistant", text),
                actual_model=self.identity.configured_model,
            )
        return ModelResponse(
            CanonicalMessage(
                "assistant",
                "",
                tool_calls=[
                    CanonicalToolCall(
                        "call-{}".format(agent_id),
                        "list_directory",
                        {"path": "."},
                    )
                ],
            ),
            actual_model=self.identity.configured_model,
        )

    def complete_stream(self, request):
        yield "两个智能体均已读取目录。"


def _make_config():
    from src.core.config.loader import (
        AgentPreset,
        AppConfig,
        Capability,
        ModelProfile,
        ProjectConfig,
        ToolConfig,
    )

    app = AppConfig(
        default_agent="general",
        timezone="Asia/Shanghai",
        history_rounds=10,
        image_prompt="",
        active_model="test_model",
        fallback_model="test_model",
        local_model="",
        flash_model="",
        pro_model="",
        vision_model="",
        embedding_model="",
        rerank_model="",
        fallback_cooldown_seconds=60,
    )
    model_profile = ModelProfile(
        id="test_model",
        enabled=True,
        type="openai_compatible",
        provider="test",
        base_url="http://localhost:11434",
        model="test-model",
        temperature=0.7,
        max_tokens=2048,
        timeout_seconds=30,
        capabilities=ModelCapabilities(tools=True, vision=False, reasoning=False),
    )
    agent = AgentPreset(
        id="general",
        name="通用助手",
        role="assistant",
        description="通用对话助手",
        system_prompt="你是一个有帮助的助手。",
        capabilities=[Capability(name="对话", description="通用对话")],
        tools=[],
    )
    tools = ToolConfig(
        enabled=False,
        default_working_directory=".",
        allowed_roots=[],
        denied_globs=[],
        approval_ttl_seconds=60,
        max_tool_rounds=5,
        max_total_tool_calls=20,
        max_read_bytes=1024,
        max_write_bytes=1024,
        max_directory_entries=100,
        max_search_results=50,
        max_command_output_bytes=4096,
        default_command_timeout_seconds=30,
        max_command_timeout_seconds=60,
        enabled_command_profiles=[],
    )
    return ProjectConfig(
        app=app,
        models={"test_model": model_profile},
        tools=tools,
        plugins={},
        agents={"general": agent},
        scripts={},
        schedules=[],
    )


class WebApiTest(unittest.TestCase):
    def setUp(self):
        from src.api.app import create_app
        import src.api.routers.chat as chat_module
        from pathlib import Path
        from src.core.services.auth import AdminAuthService
        from src.core.storage.admin_users import (
            AdminRoleStore,
            AdminSessionStore,
            AdminUserStore,
        )
        from src.core.storage.database import Database

        self.config = _make_config()
        self.fake_client = FakeClient()
        self.model_router = ModelRouter.single(self.fake_client)

        self.mock_registry = MagicMock()
        self.mock_registry.resolve.return_value = MagicMock(tenant_id="00000000-0000-0000-0000-000000000001")

        self.mock_store = MagicMock()
        self.mock_store.load_context.return_value = []

        self._db_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self._db_dir.cleanup)
        database = Database(Path(self._db_dir.name) / "botplatform.sqlite3")
        self.admin_users = AdminUserStore(database)
        self.admin_roles = AdminRoleStore(database)
        self.admin_sessions = AdminSessionStore(database, b"test-secret")
        self.admin_auth = AdminAuthService(
            self.admin_users,
            self.admin_roles,
            self.admin_sessions,
            Path(self._db_dir.name),
        )
        admin_role = self.admin_roles.get_by_code("admin")
        self.admin_users.create("admin", "password12345", admin_role.role_id)

        self.app = create_app(
            self.config, self.model_router, self.mock_registry, self.mock_store,
            admin_auth=self.admin_auth,
            admin_user_store=self.admin_users,
            admin_role_store=self.admin_roles,
        )
        self.client = TestClient(self.app)
        response = self.client.post(
            "/api/auth/login",
            json={"username": "admin", "password": "password12345"},
        )
        assert response.status_code == 200, response.text

        self._tmp = tempfile.NamedTemporaryFile(suffix=".json", delete=False)
        self._tmp.close()
        os.remove(self._tmp.name)
        self._patcher = patch.object(
            chat_module, "CONVERSATIONS_FILE", type(chat_module.CONVERSATIONS_FILE)(self._tmp.name)
        )
        self._patcher.start()
        self.addCleanup(self._patcher.stop)
        self.addCleanup(lambda: os.path.exists(self._tmp.name) and os.remove(self._tmp.name))

    def _auth_params(self):
        return {}

    def _create_conversation(self):
        response = self.client.post("/api/chat/conversations", params=self._auth_params())
        self.assertEqual(response.status_code, 201)
        return response.json()["id"]

    def test_health_no_auth(self):
        response = self.client.get("/api/health")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["status"], "ok")

    def test_status_requires_auth(self):
        from fastapi.testclient import TestClient as _TestClient

        anonymous = _TestClient(self.app)
        response = anonymous.get("/api/status")
        self.assertEqual(response.status_code, 401)

    def test_status_with_auth(self):
        response = self.client.get("/api/status", params=self._auth_params())
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["model_ready"])
        self.assertEqual(data["active_model"], "test_model")

    def test_models_list(self):
        response = self.client.get("/api/models", params=self._auth_params())
        self.assertEqual(response.status_code, 200)
        models = response.json()
        self.assertEqual(len(models), 1)
        self.assertEqual(models[0]["id"], "test_model")
        self.assertTrue(models[0]["is_primary"])

    def test_model_status(self):
        response = self.client.get("/api/models/status", params=self._auth_params())
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["primary_profile_id"], "test_model")
        self.assertFalse(data["cooling_down"])

    def test_model_not_found(self):
        response = self.client.get("/api/models/nonexistent", params=self._auth_params())
        self.assertEqual(response.status_code, 404)

    def test_agents_list(self):
        response = self.client.get("/api/agents", params=self._auth_params())
        self.assertEqual(response.status_code, 200)
        agents = response.json()
        self.assertEqual(len(agents), 1)
        self.assertEqual(agents[0]["id"], "general")

    def test_agent_active(self):
        response = self.client.get("/api/agents/active", params=self._auth_params())
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["id"], "general")

    def test_agent_not_found(self):
        response = self.client.get("/api/agents/nonexistent", params=self._auth_params())
        self.assertEqual(response.status_code, 404)

    def test_chat_stream(self):
        conv_id = self._create_conversation()
        response = self.client.post(
            "/api/chat",
            params=self._auth_params(),
            json={"message": "你好", "conversation_id": conv_id},
        )
        self.assertEqual(response.status_code, 200)
        self.assertIn("text/event-stream", response.headers["content-type"])
        body = response.text
        self.assertIn('"type": "token"', body)
        self.assertIn('"type": "done"', body)
        self.mock_store.save_context.assert_called_once()

    def test_chat_tools_use_conversation_workspace(self):
        tenant = self.mock_registry.resolve.return_value
        self.config.agents["general"].tools.append("list_directory")
        tool_runtime = MagicMock()
        tool_runtime.schemas.return_value = [
            {
                "type": "function",
                "function": {
                    "name": "list_directory",
                    "description": "列出目录",
                    "parameters": {"type": "object", "properties": {}},
                },
            }
        ]
        tool_runtime.execute.return_value = ToolResult(
            True, data={"path": ".", "entries": []}
        )
        self.app.state.tool_runtime = tool_runtime
        self.fake_client.complete = MagicMock(side_effect=[
            ModelResponse(
                CanonicalMessage(
                    "assistant",
                    "",
                    tool_calls=[
                        CanonicalToolCall(
                            "call-1", "list_directory", {"path": "."}
                        )
                    ],
                ),
                actual_model=self.fake_client.identity.configured_model,
            ),
            ModelResponse(
                CanonicalMessage("assistant", "当前目录为空。"),
                actual_model=self.fake_client.identity.configured_model,
            ),
        ])

        conv_id = self._create_conversation()
        response = self.client.post(
            "/api/chat",
            params=self._auth_params(),
            json={"message": "查看当前目录下的文件", "conversation_id": conv_id},
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn('"type": "tool_result"', response.text)
        self.assertIn('"ok": true', response.text)
        self.assertEqual(tool_runtime.bind_tenant.call_count, 2)
        tool_runtime.bind_tenant.assert_called_with(tenant)
        tool_runtime.execute.assert_called_once_with(
            "list_directory", {"path": "."}
        )

    def test_multi_agent_chat_executes_structured_directory_tools(self):
        general = replace(
            self.config.agents["general"],
            tools=["list_directory"],
        )
        coder = replace(
            general,
            id="coder",
            name="编程开发助手",
        )
        self.config.agents["general"] = general
        self.config.agents["coder"] = coder

        client = MultiAgentToolClient()
        self.model_router.clients["test_model"] = client
        tool_runtime = MagicMock()
        tool_runtime.config = self.config.tools
        tool_runtime.schemas.side_effect = lambda names: [
            {
                "type": "function",
                "function": {
                    "name": name,
                    "description": "测试工具",
                    "parameters": {"type": "object", "properties": {}},
                },
            }
            for name in names
        ]
        tool_runtime.requires_approval.return_value = False
        tool_runtime.execute.return_value = ToolResult(
            True, data={"directory": "/workspace", "items": []}
        )
        self.app.state.tool_runtime = tool_runtime

        conv_id = self._create_conversation()
        response = self.client.post(
            "/api/chat",
            params=self._auth_params(),
            json={
                "message": "查看当前目录下的文件",
                "conversation_id": conv_id,
                "agent_ids": ["general", "coder"],
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn('"type": "agent_done"', response.text)
        self.assertIn("general：目录读取完成。", response.text)
        self.assertIn("coder：目录读取完成。", response.text)
        self.assertIn("两个智能体均已读取目录。", response.text)
        self.assertNotIn("DSML", response.text)
        self.assertEqual(tool_runtime.execute.call_count, 2)

        agent_requests = [
            request
            for request in client.requests
            if request.context.operation == "agent_subtask"
        ]
        self.assertEqual(len(agent_requests), 4)
        for agent_id in ("general", "coder"):
            requests = [
                request
                for request in agent_requests
                if request.context.agent_id == agent_id
            ]
            self.assertEqual(len(requests), 2)
            self.assertEqual(
                [tool["function"]["name"] for tool in requests[0].tools],
                ["list_directory"],
            )
            roles = [message.role for message in requests[1].messages]
            self.assertEqual(roles.count("assistant"), 1)
            self.assertEqual(roles.count("tool"), 1)

    def test_multi_agent_uses_each_agents_own_tool_schema(self):
        from src.api.routers.chat import _run_agent

        tenant = self.mock_registry.resolve.return_value
        general = replace(
            self.config.agents["general"],
            tools=["list_directory"],
        )
        coder = replace(
            general,
            id="coder",
            name="编程开发助手",
            tools=["find_files"],
        )
        tool_runtime = MagicMock()
        tool_runtime.config = self.config.tools
        tool_runtime.schemas.side_effect = lambda names: [
            {
                "type": "function",
                "function": {
                    "name": name,
                    "description": "测试工具",
                    "parameters": {"type": "object", "properties": {}},
                },
            }
            for name in names
        ]
        self.fake_client.complete = MagicMock(
            return_value=ModelResponse(
                CanonicalMessage("assistant", "完成"),
                actual_model=self.fake_client.identity.configured_model,
            )
        )
        context = ModelCallContext(source="web", operation="answer")

        _run_agent(
            general, "查看目录", [], self.model_router, [], context,
            tool_runtime, tenant,
        )
        _run_agent(
            coder, "查找文件", [], self.model_router, [], context,
            tool_runtime, tenant,
        )

        requests = [
            call.args[0] for call in self.fake_client.complete.call_args_list
        ]
        self.assertEqual(
            [tool["function"]["name"] for tool in requests[0].tools],
            ["list_directory"],
        )
        self.assertEqual(
            [tool["function"]["name"] for tool in requests[1].tools],
            ["find_files"],
        )

    def test_multi_agent_blocks_tools_requiring_approval(self):
        from src.api.routers.chat import _run_agent

        tenant = self.mock_registry.resolve.return_value
        agent = replace(
            self.config.agents["general"],
            tools=["run_command"],
        )
        tool_runtime = MagicMock()
        tool_runtime.config = self.config.tools
        tool_runtime.schemas.return_value = [
            {
                "type": "function",
                "function": {
                    "name": "run_command",
                    "description": "运行命令",
                    "parameters": {"type": "object", "properties": {}},
                },
            }
        ]
        tool_runtime.requires_approval.return_value = True
        requests = []

        def complete(request):
            requests.append(request)
            if any(message.role == "tool" for message in request.messages):
                return ModelResponse(
                    CanonicalMessage("assistant", "该操作需要确认，未执行。"),
                    actual_model=self.fake_client.identity.configured_model,
                )
            return ModelResponse(
                CanonicalMessage(
                    "assistant",
                    "",
                    tool_calls=[
                        CanonicalToolCall(
                            "call-command",
                            "run_command",
                            {"command": "pwd"},
                        )
                    ],
                ),
                actual_model=self.fake_client.identity.configured_model,
            )

        self.fake_client.complete = MagicMock(side_effect=complete)
        answer = _run_agent(
            agent, "运行命令", [], self.model_router, [],
            ModelCallContext(source="web"), tool_runtime, tenant,
        )

        self.assertEqual(answer, "该操作需要确认，未执行。")
        tool_runtime.execute.assert_not_called()
        payload = json.loads(requests[-1].messages[-1].content)
        self.assertFalse(payload["ok"])
        self.assertIn("暂不执行需要确认的工具", payload["error"])

    def test_multi_agent_returns_tool_failures_to_model(self):
        from src.api.routers.chat import _run_agent

        tenant = self.mock_registry.resolve.return_value
        agent = replace(
            self.config.agents["general"],
            tools=["list_directory"],
        )
        tool_runtime = MagicMock()
        tool_runtime.config = self.config.tools
        tool_runtime.schemas.return_value = [
            {
                "type": "function",
                "function": {
                    "name": "list_directory",
                    "description": "列出目录",
                    "parameters": {"type": "object", "properties": {}},
                },
            }
        ]
        tool_runtime.requires_approval.return_value = False
        tool_runtime.execute.return_value = ToolResult(
            False, error="无法读取目录"
        )
        requests = []

        def complete(request):
            requests.append(request)
            if any(message.role == "tool" for message in request.messages):
                return ModelResponse(
                    CanonicalMessage("assistant", "目录读取失败：无法读取目录。"),
                    actual_model=self.fake_client.identity.configured_model,
                )
            return ModelResponse(
                CanonicalMessage(
                    "assistant",
                    "",
                    tool_calls=[
                        CanonicalToolCall(
                            "call-list-1", "list_directory", {"path": "."}
                        ),
                        CanonicalToolCall(
                            "call-list-2", "list_directory", {"path": "."}
                        ),
                    ],
                ),
                actual_model=self.fake_client.identity.configured_model,
            )

        self.fake_client.complete = MagicMock(side_effect=complete)
        answer = _run_agent(
            agent, "查看目录", [], self.model_router, [],
            ModelCallContext(source="web"), tool_runtime, tenant,
        )

        self.assertEqual(answer, "目录读取失败：无法读取目录。")
        payload = json.loads(requests[-1].messages[-1].content)
        self.assertEqual(payload, {"ok": False, "error": "无法读取目录"})
        roles = [message.role for message in requests[-1].messages]
        self.assertEqual(roles.count("assistant"), 1)
        self.assertEqual(roles.count("tool"), 2)
        self.assertEqual(tool_runtime.execute.call_count, 2)

    def test_multi_agent_enforces_tool_call_limits(self):
        from src.api.routers.chat import _run_agent

        tenant = self.mock_registry.resolve.return_value
        agent = replace(
            self.config.agents["general"],
            tools=["list_directory"],
        )
        schema = {
            "type": "function",
            "function": {
                "name": "list_directory",
                "description": "列出目录",
                "parameters": {"type": "object", "properties": {}},
            },
        }
        context = ModelCallContext(source="web")

        total_runtime = MagicMock()
        total_runtime.config = replace(
            self.config.tools, max_total_tool_calls=1
        )
        total_runtime.schemas.return_value = [schema]
        self.fake_client.complete = MagicMock(
            return_value=ModelResponse(
                CanonicalMessage(
                    "assistant",
                    "",
                    tool_calls=[
                        CanonicalToolCall("call-1", "list_directory", {}),
                        CanonicalToolCall("call-2", "list_directory", {}),
                    ],
                ),
                actual_model=self.fake_client.identity.configured_model,
            )
        )
        answer = _run_agent(
            agent, "查看目录", [], self.model_router, [], context,
            total_runtime, tenant,
        )
        self.assertIn("工具步骤超过安全上限", answer)
        total_runtime.execute.assert_not_called()

        round_runtime = MagicMock()
        round_runtime.config = replace(
            self.config.tools, max_tool_rounds=1
        )
        round_runtime.schemas.return_value = [schema]
        round_runtime.requires_approval.return_value = False
        round_runtime.execute.return_value = ToolResult(
            True, data={"items": []}
        )
        self.fake_client.complete = MagicMock(
            return_value=ModelResponse(
                CanonicalMessage(
                    "assistant",
                    "",
                    tool_calls=[
                        CanonicalToolCall("call-1", "list_directory", {})
                    ],
                ),
                actual_model=self.fake_client.identity.configured_model,
            )
        )
        answer = _run_agent(
            agent, "查看目录", [], self.model_router, [], context,
            round_runtime, tenant,
        )
        self.assertIn("工具调用轮次上限", answer)
        self.assertEqual(round_runtime.execute.call_count, 1)

    def test_chat_requires_conversation(self):
        response = self.client.post(
            "/api/chat",
            params=self._auth_params(),
            json={"message": "你好"},
        )
        self.assertEqual(response.status_code, 400)

    def test_chat_history(self):
        conv_id = self._create_conversation()
        self.mock_store.load_context.return_value = [
            CanonicalMessage("user", "你好"),
            CanonicalMessage("assistant", "你好！"),
        ]
        response = self.client.get(
            "/api/chat/history",
            params=dict(self._auth_params(), conversation_id=conv_id),
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(len(data["messages"]), 2)

    def test_conversation_list_and_delete(self):
        conv_id = self._create_conversation()
        response = self.client.get("/api/chat/conversations", params=self._auth_params())
        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(response.json()), 1)

        response = self.client.delete(
            "/api/chat/conversations/" + conv_id, params=self._auth_params()
        )
        self.assertEqual(response.status_code, 200)

        response = self.client.get("/api/chat/conversations", params=self._auth_params())
        self.assertEqual(len(response.json()), 0)

    def test_page_chat(self):
        response = self.client.get("/", params=self._auth_params())
        self.assertEqual(response.status_code, 200)
        self.assertIn("chat-messages", response.text)

    def test_chat_uses_local_markdown_dependencies_under_csp(self):
        response = self.client.get("/", params=self._auth_params())
        self.assertEqual(response.status_code, 200)
        self.assertNotIn("cdn.jsdelivr.net", response.text)
        marked = response.text.index("/static/vendor/marked/marked.umd.js")
        purify = response.text.index("/static/vendor/dompurify/purify.min.js")
        chat = response.text.index("/static/js/chat.js")
        self.assertLess(marked, purify)
        self.assertLess(purify, chat)
        for path, marker in (
            ("/static/vendor/marked/marked.umd.js", "marked v18.0.7"),
            ("/static/vendor/dompurify/purify.min.js", "DOMPurify 3.4.12"),
        ):
            asset = self.client.get(path)
            self.assertEqual(asset.status_code, 200)
            self.assertIn(marker, asset.text)

    def test_reworked_management_pages_render(self):
        scripts = self.client.get("/scripts")
        self.assertEqual(scripts.status_code, 200)
        self.assertIn('data-scripts-tab="catalog"', scripts.text)
        self.assertIn('data-scripts-tab="runs"', scripts.text)
        self.assertIn("script-settings-modal", scripts.text)
        self.assertIn('href="/scripts" class="nav-sub-item active"', scripts.text)
        self.assertIn('nav-item nav-group-toggle active', scripts.text)

        tools = self.client.get("/tools")
        self.assertEqual(tools.status_code, 200)
        self.assertIn("执行审计", tools.text)
        self.assertIn("audit-filter-status", tools.text)
        tool_menu_labels = [
            "内置工具",
            "Skill 技能",
            "MCP 服务",
            "运维脚本",
            "执行审计",
        ]
        tool_menu_positions = [
            tools.text.index(">{}<".format(label)) for label in tool_menu_labels
        ]
        self.assertEqual(tool_menu_positions, sorted(tool_menu_positions))
        section_positions = [
            tools.text.index(
                '<div class="nav-section-label">{}</div>'.format(label)
            )
            for label in ("工作台", "智能能力", "内容资源", "运营", "系统")
        ]
        self.assertEqual(section_positions, sorted(section_positions))

        users = self.client.get("/users")
        self.assertEqual(users.status_code, 200)
        self.assertIn("用户与权限 - BotPlatform", users.text)
        self.assertIn("tenant-detail-modal", users.text)
        self.assertIn("tenant-detail-title", users.text)

    def test_management_layout_structure(self):
        tools = self.client.get("/tools")
        self.assertEqual(tools.status_code, 200)
        tools_dom = BeautifulSoup(tools.text, "html.parser")
        category_tabs = tools_dom.select_one("#builtin-category-tabs")
        self.assertIsNotNone(category_tabs)
        self.assertEqual(category_tabs.get("role"), "tablist")
        builtin_list = tools_dom.select_one("#builtin-tools-list")
        self.assertIsNotNone(builtin_list)
        self.assertEqual(builtin_list.get("role"), "tabpanel")

        sidebar_actions = tools_dom.select_one(".sidebar-user > .sidebar-user-actions")
        self.assertIsNotNone(sidebar_actions)
        self.assertIsNotNone(sidebar_actions.select_one("#theme-toggle[aria-label='切换主题']"))
        self.assertIsNotNone(sidebar_actions.select_one("#logout-btn[aria-label='退出登录']"))

        drive = self.client.get("/drive")
        self.assertEqual(drive.status_code, 200)
        drive_dom = BeautifulSoup(drive.text, "html.parser")
        files_panel = drive_dom.select_one("#drive-files-panel")
        self.assertIsNotNone(files_panel.select_one(".drive-panel-toolbar .drive-toolbar"))
        self.assertIsNotNone(files_panel.select_one("#drive-mkdir-btn"))
        self.assertIsNotNone(files_panel.select_one("#drive-newfile-btn"))
        self.assertIsNotNone(files_panel.select_one("#drive-upload-btn"))
        self.assertIsNone(drive_dom.select_one(".page-header .drive-toolbar"))
        self.assertIsNone(drive_dom.select_one("#drive-audit-panel .drive-toolbar"))

    def test_page_models(self):
        response = self.client.get("/models", params=self._auth_params())
        self.assertEqual(response.status_code, 200)
        self.assertIn("model-list", response.text)

    def test_page_agents(self):
        response = self.client.get("/agents", params=self._auth_params())
        self.assertEqual(response.status_code, 200)
        self.assertIn("agent-list", response.text)

    def test_login_sets_session_cookie(self):
        from fastapi.testclient import TestClient as _TestClient

        client = _TestClient(self.app)
        response = client.post(
            "/api/auth/login",
            json={"username": "admin", "password": "password12345"},
        )
        self.assertEqual(response.status_code, 200)
        self.assertIn("admin_session", response.cookies)
        self.assertEqual(response.json()["user"]["username"], "admin")

    def test_unauthenticated_page_redirects_to_login(self):
        from fastapi.testclient import TestClient as _TestClient

        anonymous = _TestClient(self.app)
        response = anonymous.get("/models", follow_redirects=False)
        self.assertEqual(response.status_code, 302)
        self.assertTrue(response.headers["location"].startswith("/login"))

    def test_logout_invalidates_session(self):
        response = self.client.post("/api/auth/logout")
        self.assertEqual(response.status_code, 200)
        response = self.client.get("/api/status")
        self.assertEqual(response.status_code, 401)

    def test_security_headers_include_csp(self):
        response = self.client.get("/api/health")
        csp = response.headers.get("Content-Security-Policy", "")
        self.assertIn("default-src 'self'", csp)
        self.assertIn("script-src 'self' 'unsafe-inline'", csp)
        self.assertIn("frame-ancestors 'none'", csp)

    def test_session_cookie_secure_flag_follows_app_setting(self):
        from fastapi.testclient import TestClient as _TestClient
        from src.api.app import create_app

        # Default app: the session cookie must not carry the Secure attribute.
        response = _TestClient(self.app).post(
            "/api/auth/login",
            json={"username": "admin", "password": "password12345"},
        )
        self.assertEqual(response.status_code, 200)
        self.assertNotIn("secure", response.headers["set-cookie"].lower())

        secure_app = create_app(
            self.config, self.model_router, self.mock_registry, self.mock_store,
            admin_auth=self.admin_auth,
            admin_user_store=self.admin_users,
            admin_role_store=self.admin_roles,
            secure_cookies=True,
        )
        response = _TestClient(secure_app).post(
            "/api/auth/login",
            json={"username": "admin", "password": "password12345"},
        )
        self.assertEqual(response.status_code, 200)
        self.assertIn("secure", response.headers["set-cookie"].lower())

    def test_login_lockout_returns_429_after_repeated_failures(self):
        from fastapi.testclient import TestClient as _TestClient

        anonymous = _TestClient(self.app)
        for _ in range(5):
            response = anonymous.post(
                "/api/auth/login",
                json={"username": "admin", "password": "wrong-password"},
            )
            self.assertEqual(response.status_code, 401)
        # Even the correct password is rejected while the pair is locked out.
        response = anonymous.post(
            "/api/auth/login",
            json={"username": "admin", "password": "password12345"},
        )
        self.assertEqual(response.status_code, 429)


class LoginThrottleTest(unittest.TestCase):
    """Service-level throttle behaviour with an injectable clock."""

    def setUp(self):
        from pathlib import Path
        from src.core.services.auth import AdminAuthService
        from src.core.storage.admin_users import (
            AdminRoleStore,
            AdminSessionStore,
            AdminUserStore,
        )
        from src.core.storage.database import Database

        self._dir = tempfile.TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        database = Database(Path(self._dir.name) / "botplatform.sqlite3")
        users = AdminUserStore(database)
        roles = AdminRoleStore(database)
        sessions = AdminSessionStore(database, b"test-secret")
        self.now = 1000.0
        self.auth = AdminAuthService(
            users, roles, sessions, Path(self._dir.name), monotonic=lambda: self.now
        )
        role = roles.get_by_code("admin")
        users.create("admin", "password12345", role.role_id)

    def _fail(self, ip="1.2.3.4"):
        from src.core.services.auth import AuthError

        with self.assertRaises(AuthError):
            self.auth.login("admin", "wrong-password", ip=ip)

    def test_lockout_scope_and_expiry(self):
        from src.core.services.auth import (
            LOGIN_LOCKOUT_SECONDS,
            MAX_LOGIN_FAILURES,
            LoginThrottled,
        )

        for _ in range(MAX_LOGIN_FAILURES):
            self._fail()
        with self.assertRaises(LoginThrottled):
            self.auth.login("admin", "password12345", ip="1.2.3.4")
        # A different source IP is not affected by the lockout.
        token, principal = self.auth.login("admin", "password12345", ip="5.6.7.8")
        self.assertTrue(token)
        self.assertEqual(principal.user.username, "admin")
        # After the lockout expires the locked pair may log in again.
        self.now += LOGIN_LOCKOUT_SECONDS + 1
        token, _ = self.auth.login("admin", "password12345", ip="1.2.3.4")
        self.assertTrue(token)

    def test_successful_login_resets_failure_counter(self):
        for _ in range(4):
            self._fail()
        self.auth.login("admin", "password12345", ip="1.2.3.4")
        for _ in range(4):
            self._fail()
        # The counter was reset, so the pair is still below the limit.
        token, _ = self.auth.login("admin", "password12345", ip="1.2.3.4")
        self.assertTrue(token)


class KnowledgeApiTest(unittest.TestCase):
    """Knowledge management endpoints backed by a real tenant registry."""

    def setUp(self):
        from pathlib import Path
        from src.api.app import create_app
        from src.core.services.auth import AdminAuthService
        from src.core.services.knowledge import KnowledgeService
        from src.core.storage.admin_users import (
            AdminRoleStore,
            AdminSessionStore,
            AdminUserStore,
        )
        from src.core.storage.database import Database
        from src.core.storage.tenants import TenantRegistry

        self._dir = tempfile.TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        root = Path(self._dir.name)
        self.registry = TenantRegistry(root / "data")
        self.tenant = self.registry.resolve("bot", "user")
        self.knowledge = KnowledgeService(self.registry, None)

        database = Database(root / "admin.sqlite3")
        users = AdminUserStore(database)
        roles = AdminRoleStore(database)
        sessions = AdminSessionStore(database, b"test-secret")
        admin_auth = AdminAuthService(users, roles, sessions, root)
        role = roles.get_by_code("admin")
        users.create("admin", "password12345", role.role_id)

        self.app = create_app(
            _make_config(), ModelRouter.single(FakeClient()),
            self.registry, MagicMock(),
            knowledge_service=self.knowledge,
            admin_auth=admin_auth,
            admin_user_store=users,
            admin_role_store=roles,
        )
        self.client = TestClient(self.app)
        response = self.client.post(
            "/api/auth/login",
            json={"username": "admin", "password": "password12345"},
        )
        assert response.status_code == 200, response.text

    def test_endpoints_require_auth(self):
        anonymous = TestClient(self.app)
        response = anonymous.get("/api/knowledge/tenants")
        self.assertEqual(response.status_code, 401)
        response = anonymous.get(
            "/api/knowledge", params={"tenant_id": self.tenant.tenant_id}
        )
        self.assertEqual(response.status_code, 401)

    def test_tenant_dropdown_lists_contexts(self):
        response = self.client.get("/api/knowledge/tenants")
        self.assertEqual(response.status_code, 200)
        tenants = response.json()
        self.assertEqual(len(tenants), 1)
        self.assertEqual(tenants[0]["tenant_id"], self.tenant.tenant_id)
        self.assertEqual(tenants[0]["bot_id"], "bot")
        self.assertEqual(tenants[0]["user_id"], "user")

    def test_upload_list_search_delete_flow(self):
        response = self.client.post(
            "/api/knowledge/upload",
            data={"tenant_id": self.tenant.tenant_id},
            files={"file": (
                "faq.md",
                "# FAQ\n\n退货政策支持七天无理由退货。".encode("utf-8"),
                "text/markdown",
            )},
        )
        self.assertEqual(response.status_code, 200, response.text)
        self.assertGreater(response.json()["chunks"], 0)
        saved = (
            self.registry.tenant_root(self.tenant.tenant_id)
            / "workspace" / "knowledge_uploads" / "faq.md"
        )
        self.assertTrue(saved.is_file())

        listing = self.client.get(
            "/api/knowledge", params={"tenant_id": self.tenant.tenant_id}
        )
        self.assertEqual(listing.status_code, 200)
        sources = listing.json()["sources"]
        self.assertEqual(len(sources), 1)
        self.assertEqual(sources[0]["source_type"], "file")

        search = self.client.get(
            "/api/knowledge/search",
            params={"tenant_id": self.tenant.tenant_id, "q": "退货政策"},
        )
        self.assertEqual(search.status_code, 200)
        results = search.json()["results"]
        self.assertTrue(any("退货" in item["content"] for item in results))

        deleted = self.client.delete(
            "/api/knowledge/" + sources[0]["source_id"],
            params={"tenant_id": self.tenant.tenant_id},
        )
        self.assertEqual(deleted.status_code, 200)
        listing = self.client.get(
            "/api/knowledge", params={"tenant_id": self.tenant.tenant_id}
        )
        self.assertEqual(listing.json()["sources"], [])

    def test_upload_docx_document_is_parsed_and_searchable(self):
        import io
        import docx

        buffer = io.BytesIO()
        document = docx.Document()
        document.add_heading("产品手册", level=1)
        document.add_paragraph("保修期为一年。")
        document.save(buffer)
        response = self.client.post(
            "/api/knowledge/upload",
            data={"tenant_id": self.tenant.tenant_id},
            files={"file": ("manual.docx", buffer.getvalue(), "application/octet-stream")},
        )
        self.assertEqual(response.status_code, 200, response.text)
        self.assertGreater(response.json()["chunks"], 0)
        search = self.client.get(
            "/api/knowledge/search",
            params={"tenant_id": self.tenant.tenant_id, "q": "保修期"},
        )
        self.assertTrue(
            any("保修期" in item["content"] for item in search.json()["results"])
        )

    def test_upload_rejects_unsupported_suffix_and_broken_document(self):
        response = self.client.post(
            "/api/knowledge/upload",
            data={"tenant_id": self.tenant.tenant_id},
            files={"file": ("payload.exe", b"MZ", "application/octet-stream")},
        )
        self.assertEqual(response.status_code, 400)

        response = self.client.post(
            "/api/knowledge/upload",
            data={"tenant_id": self.tenant.tenant_id},
            files={"file": ("broken.docx", b"not a real docx", "application/octet-stream")},
        )
        self.assertEqual(response.status_code, 400)
        # The broken upload must not linger on disk.
        leftover = (
            self.registry.tenant_root(self.tenant.tenant_id)
            / "workspace" / "knowledge_uploads" / "broken.docx"
        )
        self.assertFalse(leftover.exists())

    def test_add_text_and_reindex_without_embedding(self):
        response = self.client.post(
            "/api/knowledge/text",
            json={
                "tenant_id": self.tenant.tenant_id,
                "name": "产品FAQ",
                "content": "客服工作时间为工作日九点到六点。",
            },
        )
        self.assertEqual(response.status_code, 200, response.text)
        self.assertEqual(response.json()["status"], "pending_embedding")

        response = self.client.post(
            "/api/knowledge/reindex",
            json={"tenant_id": self.tenant.tenant_id},
        )
        self.assertEqual(response.status_code, 400)
        self.assertIn("embedding", response.json()["detail"])

    def test_unknown_tenant_returns_404(self):
        import uuid

        response = self.client.get(
            "/api/knowledge", params={"tenant_id": str(uuid.uuid4())}
        )
        self.assertEqual(response.status_code, 404)


if __name__ == "__main__":
    unittest.main()
